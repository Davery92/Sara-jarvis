"""
Document Renderer

Turns a markdown document body into a real downloadable file — Word (.docx)
today, PDF once WeasyPrint + its system deps land (Part A-3). Markdown is the
single input format: Sara is reliable at markdown and one input keeps the
`document_generate` tool schema trivial.

The docx path uses python-docx (already installed, no Dockerfile change). The
pdf path is intentionally isolated behind `render_pdf` so A-3 can wire in
markdown->HTML->WeasyPrint without touching the docx code or the tool.
"""

from __future__ import annotations

import re
import logging
from io import BytesIO
from typing import List, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Inline markdown (bold / italic / code / links) -> python-docx runs
# ---------------------------------------------------------------------------

# Order matters: code spans first so we don't parse markup inside them.
_INLINE_RE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<bold>\*\*[^*]+\*\*|__[^_]+__)"
    r"|(?P<italic>\*[^*]+\*|_[^_]+_)"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Append runs to a paragraph, honoring inline markdown formatting."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("code"):
            run = paragraph.add_run(m.group("code")[1:-1])
            run.font.name = "Consolas"
        elif m.group("bold"):
            run = paragraph.add_run(m.group("bold")[2:-2])
            run.bold = True
        elif m.group("italic"):
            run = paragraph.add_run(m.group("italic")[1:-1])
            run.italic = True
        elif m.group("link"):
            link = m.group("link")
            label = link[1:link.index("]")]
            url = link[link.index("(") + 1:-1]
            # python-docx has no simple hyperlink helper; show "label (url)".
            run = paragraph.add_run(label if label == url else f"{label} ({url})")
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_table_divider(line: str) -> bool:
    s = line.strip().strip("|")
    return bool(s) and all(re.fullmatch(r"\s*:?-+:?\s*", c) for c in s.split("|"))


def _split_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


# ---------------------------------------------------------------------------
# docx
# ---------------------------------------------------------------------------

def render_docx(title: str, markdown_body: str, style: str = "default") -> bytes:
    """Render markdown to a .docx byte string."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title block. 'letter'/'report' styles get a centered title; 'report' adds
    # a page break so the body starts fresh (a lightweight title page).
    if title:
        heading = doc.add_heading(title, level=0)
        if style in ("letter", "report"):
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if style == "report":
            doc.add_page_break()

    lines = markdown_body.replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines: List[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            for cl in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(cl if cl else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue

        # Table
        if _is_table_row(line) and i + 1 < n and _is_table_divider(lines[i + 1]):
            header = _split_row(line)
            i += 2  # skip header + divider
            body_rows: List[List[str]] = []
            while i < n and _is_table_row(lines[i]):
                body_rows.append(_split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            try:
                table.style = "Light Grid Accent 1"
            except Exception:
                pass
            for idx, cell_text in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.paragraphs[0].add_run(cell_text).bold = True
            for row in body_rows:
                cells = table.add_row().cells
                for idx in range(len(header)):
                    val = row[idx] if idx < len(row) else ""
                    _add_inline_runs(cells[idx].paragraphs[0], val)
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"(\*\s*){3,}|(-\s*){3,}|(_\s*){3,}", stripped):
            doc.add_paragraph("―" * 30)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote" if _has_style(doc, "Intense Quote") else None)
            _add_inline_runs(p, stripped.lstrip(">").strip())
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*+]\s+", stripped):
            _add_inline_runs(
                doc.add_paragraph(style="List Bullet"),
                re.sub(r"^[-*+]\s+", "", stripped),
            )
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+[.)]\s+", stripped):
            _add_inline_runs(
                doc.add_paragraph(style="List Number"),
                re.sub(r"^\d+[.)]\s+", "", stripped),
            )
            i += 1
            continue

        # Plain paragraph
        _add_inline_runs(doc.add_paragraph(), stripped)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _has_style(doc, name: str) -> bool:
    try:
        _ = doc.styles[name]
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# pdf (wired in A-3 — WeasyPrint + system deps + Dockerfile rebuild)
# ---------------------------------------------------------------------------

def render_pdf(title: str, markdown_body: str, style: str = "default") -> bytes:
    """Render markdown to a PDF byte string via markdown -> HTML -> WeasyPrint."""
    try:
        import markdown as md
        from weasyprint import HTML
    except ImportError as e:  # pragma: no cover - depends on A-3 rebuild
        raise RuntimeError(
            "PDF generation needs WeasyPrint + python-markdown, which aren't "
            "installed yet (Part A-3 adds them to requirements + the Dockerfile). "
            "Word (.docx) generation works now."
        ) from e

    html_body = md.markdown(
        markdown_body,
        extensions=["extra", "sane_lists", "tables", "nl2br"],
    )
    css = _PDF_STYLES.get(style, _PDF_STYLES["default"])
    title_html = f"<h1 class='doc-title'>{_escape(title)}</h1>" if title else ""
    document = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{css}</style></head>
<body>{title_html}{html_body}</body></html>"""
    return HTML(string=document).write_pdf()


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


_PDF_BASE = """
@page { margin: 2.5cm; }
body { font-family: Georgia, 'Times New Roman', serif; font-size: 11pt;
       line-height: 1.5; color: #1a1a1a; }
h1, h2, h3, h4 { font-family: Georgia, serif; line-height: 1.25; }
h1.doc-title { text-align: center; font-size: 24pt; margin-bottom: 1.5em; }
code, pre { font-family: 'Consolas', monospace; font-size: 9.5pt;
            background: #f4f4f4; }
pre { padding: 0.75em; border-radius: 4px; overflow-x: auto; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #ccc; padding: 6px 10px; text-align: left; }
th { background: #f0f0f0; }
blockquote { border-left: 3px solid #ccc; margin-left: 0; padding-left: 1em;
             color: #555; font-style: italic; }
"""

_PDF_STYLES = {
    "default": _PDF_BASE,
    "letter": _PDF_BASE + "\nbody { font-size: 12pt; } h1.doc-title { text-align: left; }",
    "report": _PDF_BASE + "\n@page { @bottom-center { content: counter(page); } }",
}


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------

_FORMATS = {
    "docx": (
        render_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    "pdf": (render_pdf, "application/pdf"),
}


def render_document(
    fmt: str, title: str, markdown_body: str, style: str = "default"
) -> Tuple[bytes, str, str]:
    """
    Render a document.

    Returns (file_bytes, filename, mime_type).
    """
    fmt = (fmt or "").lower()
    if fmt not in _FORMATS:
        raise ValueError(f"Unsupported format '{fmt}'. Use one of: {', '.join(_FORMATS)}")

    renderer, mime = _FORMATS[fmt]
    file_bytes = renderer(title, markdown_body, style)
    safe_title = re.sub(r"[^\w\- ]+", "", title or "document").strip() or "document"
    safe_title = re.sub(r"\s+", "_", safe_title)[:80]
    filename = f"{safe_title}.{fmt}"
    return file_bytes, filename, mime
